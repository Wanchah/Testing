"""
EduMorph AI Services Blueprint
Handles document processing, AI integration, and content generation.

This module provides:
- Document upload and processing (PDF, DOCX, TXT, PPT)
- YouTube video content extraction
- AI-powered note generation
- Flashcard and quiz creation
- AI chat for student questions
- Content storage and retrieval
"""

import os
import requests
import json
import re
from datetime import datetime
from flask import Blueprint, request, jsonify, render_template, current_app, flash, redirect, url_for
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from database.models import Document, Content, Flashcard, Question, AIChat, Lesson, AgeGroup, ContentFormat
from app import db
import PyPDF2
from docx import Document as DocxDocument
import openai
try:
    from openai import OpenAI
except Exception:
    OpenAI = None
from pytube import YouTube
import yt_dlp

ai_bp = Blueprint('ai_services', __name__, url_prefix='/ai')

# Configure AI services with fallbacks
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')
HUGGINGFACE_API_KEY = os.environ.get('HUGGINGFACE_API_KEY', '')
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
SERPAPI_KEY = os.environ.get('SERPAPI_KEY', '')

# Initialize OpenAI legacy and modern clients if possible
openai_client = None
if OPENAI_API_KEY:
    try:
        openai.api_key = OPENAI_API_KEY
        if OpenAI is not None:
            openai_client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception:
        pass

# Supported file types
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'ppt', 'pptx', 'md'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@ai_bp.route('/upload-document', methods=['GET', 'POST'])
@login_required
def upload_document():
    """Upload and process documents for AI content generation."""
    
    if request.method == 'POST':
        if 'document' not in request.files:
            flash('No file selected', 'error')
            return redirect(request.url)
        
        file = request.files['document']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            subject = request.form.get('subject', 'general')
            
            # Save file temporarily
            temp_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            file.save(temp_path)
            
            try:
                # Extract text content
                content = extract_document_content(temp_path, filename)
                
                # Generate AI content
                ai_content = generate_ai_content(content, subject)
                
                # Save to database
                document = Document(
                    filename=filename,
                    subject=subject,
                    user_id=current_user.id,
                    content_length=len(content),
                    file_type=filename.rsplit('.', 1)[1].lower()
                )
                db.session.add(document)
                db.session.flush()
                
                # Save extracted content
                content_obj = Content(
                    document_id=document.id,
                    raw_content=content,
                    ai_generated_notes=ai_content['notes'],
                    ai_generated_summary=ai_content['summary'],
                    key_concepts=ai_content['key_concepts'],
                    user_id=current_user.id
                )
                db.session.add(content_obj)
                
                # Generate flashcards
                flashcards = generate_flashcards(content, ai_content['key_concepts'])
                for flashcard_data in flashcards:
                    flashcard = Flashcard(
                        term=flashcard_data['term'],
                        definition=flashcard_data['definition'],
                        content_id=content_obj.id,
                        ai_generated=True
                    )
                    db.session.add(flashcard)
                
                # Generate questions
                questions = generate_questions(content, ai_content['key_concepts'])
                for question_data in questions:
                    question = Question(
                        question_text=question_data['question'],
                        answer_text=question_data['answer'],
                        question_type=question_data['type'],
                        content_id=content_obj.id,
                        ai_generated=True
                    )
                    db.session.add(question)
                
                # Also create a Lesson from this content
                try:
                    lesson = create_lesson_from_content(
                        title=os.path.splitext(filename)[0],
                        subject=subject,
                        summary=ai_content['summary'],
                        user_id=current_user.id,
                        age_group=current_user.age_group if hasattr(current_user, 'age_group') else AgeGroup.TEENS,
                        format_type=filename.rsplit('.', 1)[1].lower()
                    )
                    db.session.add(lesson)
                except Exception as _:
                    pass

                db.session.commit()
                
                # Clean up temp file
                os.remove(temp_path)
                
                flash('Document processed successfully! Lesson created from upload.', 'success')
                return redirect(url_for('ai_services.view_content', content_id=content_obj.id))
                
            except Exception as e:
                db.session.rollback()
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                flash(f'Error processing document: {str(e)}', 'error')
                return redirect(request.url)
    
    return render_template('ai_services/upload_document.html')



@ai_bp.route('/ai-chat', methods=['GET', 'POST'])
@login_required
def ai_chat():
    """AI chat interface for student questions and learning support."""
    
    if request.method == 'POST':
        message = request.form.get('message', '').strip()
        context = request.form.get('context', 'general')
        
        if not message:
            flash('Please enter a message', 'error')
            return redirect(request.url)
        
        try:
            # Get AI response
            ai_response = get_ai_response(message, context, current_user)
            
            # Save chat to database
            chat = AIChat(
                user_id=current_user.id,
                user_message=message,
                ai_response=ai_response,
                context=context,
                timestamp=datetime.utcnow()
            )
            db.session.add(chat)
            db.session.commit()
            
            flash('AI response generated successfully!', 'success')
            return redirect(url_for('ai_services.chat_history'))
            
        except Exception as e:
            flash(f'Error getting AI response: {str(e)}', 'error')
            return redirect(request.url)
    
    return render_template('ai_services/ai_chat.html')

@ai_bp.route('/chat-history')
@login_required
def chat_history():
    """View AI chat history for the current user."""
    
    chats = AIChat.query.filter_by(user_id=current_user.id).order_by(AIChat.timestamp.desc()).limit(50).all()
    return render_template('ai_services/chat_history.html', chats=chats)



# Helper functions

def extract_document_content(file_path, filename):
    """Extract text content from various document formats."""
    
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    if file_ext == 'pdf':
        return extract_pdf_content(file_path)
    elif file_ext == 'docx':
        return extract_docx_content(file_path)
    elif file_ext in ['txt', 'md']:
        return extract_text_content(file_path)
    elif file_ext in ['ppt', 'pptx']:
        return extract_powerpoint_content(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

def extract_pdf_content(file_path):
    """Extract text from PDF files."""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting PDF content: {str(e)}")

def extract_docx_content(file_path):
    """Extract text from DOCX files."""
    try:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting DOCX content: {str(e)}")

def extract_text_content(file_path):
    """Extract text from plain text files."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        raise Exception(f"Error extracting text content: {str(e)}")

def extract_powerpoint_content(file_path):
    """Extract text from PowerPoint files."""
    # This is a simplified version - you might want to use python-pptx for better extraction
    try:
        # For now, return a placeholder
        return "PowerPoint content extraction not yet implemented. Please convert to PDF or DOCX."
    except Exception as e:
        raise Exception(f"Error extracting PowerPoint content: {str(e)}")

def extract_youtube_content(url):
    """Extract content from YouTube videos."""
    try:
        # Use yt-dlp for better compatibility
        ydl_opts = {
            'writesubtitles': True,
            'writeautomaticsub': True,
            'subtitleslangs': ['en'],
            'skip_download': True
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            
            # Try to get subtitles
            transcript = ""
            if 'subtitles' in info and 'en' in info['subtitles']:
                subtitle_url = info['subtitles']['en'][0]['url']
                response = requests.get(subtitle_url)
                if response.status_code == 200:
                    transcript = response.text
            elif 'automatic_captions' in info and 'en' in info['automatic_captions']:
                subtitle_url = info['automatic_captions']['en'][0]['url']
                response = requests.get(subtitle_url)
                if response.status_code == 200:
                    transcript = response.text
            
            # If no subtitles, use video description and title
            if not transcript:
                transcript = f"Title: {info.get('title', '')}\n\nDescription: {info.get('description', '')}"
            
            return {
                'title': info.get('title', 'Unknown'),
                'transcript': transcript,
                'duration': info.get('duration', 0),
                'channel': info.get('uploader', 'Unknown'),
                'views': info.get('view_count', 0)
            }
            
    except Exception as e:
        raise Exception(f"Error extracting YouTube content: {str(e)}")

def generate_ai_content(text, subject):
    """Generate AI-powered content from text with multiple fallbacks."""
    
    try:
        # Try OpenAI first if API key is available
        if OPENAI_API_KEY and openai.api_key:
            return generate_openai_content(text, subject)
        elif HUGGINGFACE_API_KEY:
            return generate_huggingface_content(text, subject)
        else:
            # Fallback to enhanced basic content generation
            return generate_enhanced_basic_content(text, subject)
            
    except Exception as e:
        print(f"AI generation error: {str(e)}")
        # Always fallback to basic content generation
        return generate_enhanced_basic_content(text, subject)

def generate_openai_content(text, subject):
    """Generate content using OpenAI API with optimizations."""
    
    try:
        # Optimize text length for faster processing
        max_text_length = 2000
        if len(text) > max_text_length:
            text = text[:max_text_length] + "..."

        combined_prompt = (
            f"Subject: {subject}\nContent: {text}\n\n"
            "Provide a structured response with:\n"
            "1. SUMMARY: Brief 2-3 sentence summary\n"
            "2. KEY_CONCEPTS: 3-5 main concepts (one per line)\n"
            "3. NOTES: Structured educational notes\n"
            "Keep responses concise and educational."
        )

        ai_response = None
        # Prefer modern client if available
        if openai_client is not None:
            try:
                resp = openai_client.responses.create(
                    model="gpt-4o-mini",
                    input=[
                        {
                            "role": "system",
                            "content": f"You are an expert educator in {subject}. Provide structured, concise responses."
                        },
                        {"role": "user", "content": combined_prompt}
                    ],
                    max_output_tokens=400,
                )
                ai_response = getattr(resp, "output_text", None) or str(resp)
            except Exception:
                ai_response = None

        if not ai_response:
            # Fallback to legacy ChatCompletion
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": f"You are an expert educator in {subject}. Provide structured, concise responses."},
                    {"role": "user", "content": combined_prompt}
                ],
                max_tokens=400,
                temperature=0.5,
                timeout=15
            )
            ai_response = response.choices[0].message.content
        
        # Parse the combined response
        lines = ai_response.split('\n')
        summary = ""
        key_concepts = []
        notes = ai_response
        
        current_section = None
        for line in lines:
            line = line.strip()
            if line.startswith('SUMMARY:'):
                current_section = 'summary'
                summary = line.replace('SUMMARY:', '').strip()
            elif line.startswith('KEY_CONCEPTS:'):
                current_section = 'concepts'
                concepts_text = line.replace('KEY_CONCEPTS:', '').strip()
                if concepts_text:
                    key_concepts.append(concepts_text)
            elif line.startswith('NOTES:'):
                current_section = 'notes'
                notes = line.replace('NOTES:', '').strip()
            elif current_section == 'concepts' and line:
                key_concepts.append(line)
        
        # Fallback if parsing fails
        if not summary:
            summary = ai_response[:200] + "..." if len(ai_response) > 200 else ai_response
        if not key_concepts:
            key_concepts = [f'{subject.title()} Concept 1', f'{subject.title()} Concept 2', f'{subject.title()} Concept 3']
        
        return {
            'summary': summary,
            'key_concepts': key_concepts[:5],  # Limit to 5 concepts
            'notes': notes
        }
        
    except Exception as e:
        raise Exception(f"OpenAI API error: {str(e)}")

def generate_huggingface_content(text, subject):
    """Generate content using Hugging Face API."""
    # Implementation for Hugging Face API
    pass

def generate_basic_content(text, subject):
    """Generate basic content without AI APIs."""
    
    # Simple text processing
    sentences = text.split('.')
    summary = '. '.join(sentences[:3]) + '.'
    
    # Extract key concepts (simple approach)
    words = text.lower().split()
    word_freq = {}
    for word in words:
        if len(word) > 4 and word.isalpha():
            word_freq[word] = word_freq.get(word, 0) + 1
    
    key_concepts = ', '.join(sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:8])
    
    # Basic notes
    notes = f"Subject: {subject}\n\nKey Points:\n{summary}\n\nImportant Terms: {key_concepts}"
    
    return {
        'summary': summary,
        'key_concepts': key_concepts,
        'notes': notes
    }

def generate_enhanced_basic_content(text, subject):
    """Generate enhanced basic content without AI APIs."""
    
    # Enhanced text processing
    words = text.split()
    sentences = text.split('.')
    
    # Extract key concepts (words longer than 5 characters, capitalized)
    key_words = [word.strip('.,!?;:') for word in words if len(word.strip('.,!?;:')) > 5 and word[0].isupper()][:8]
    
    # Create a more meaningful summary
    first_sentences = [s.strip() for s in sentences[:3] if len(s.strip()) > 20]
    summary_text = '. '.join(first_sentences) if first_sentences else f"This content covers important {subject} concepts and provides educational value."
    
    # Generate structured notes
    notes = f"📚 {subject.title()} Study Notes\n\n"
    notes += f"📖 Summary:\n{summary_text}\n\n"
    notes += f"🔑 Key Concepts:\n" + "\n".join([f"• {word}" for word in key_words[:5]]) + "\n\n"
    notes += f"📝 Important Points:\n• This content contains valuable information about {subject}\n• Review the key concepts for better understanding\n• Practice with the generated flashcards and questions"
    
    return {
        'notes': notes,
        'summary': summary_text,
        'key_concepts': key_words[:5] if key_words else [f"{subject.title()} Concept {i+1}" for i in range(3)]
    }

def generate_flashcards(text, key_concepts):
    """Generate flashcards from content."""
    
    try:
        if openai.api_key:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Generate 5 educational flashcards in JSON format with 'term' and 'definition' fields."},
                    {"role": "user", "content": f"Content: {text[:2000]}\nKey concepts: {key_concepts}"}
                ],
                max_tokens=300
            )
            
            try:
                flashcards_data = json.loads(response.choices[0].message.content)
                return flashcards_data if isinstance(flashcards_data, list) else []
            except:
                pass
        
        # Fallback to basic flashcard generation
        return generate_basic_flashcards(text, key_concepts)
        
    except Exception as e:
        return generate_basic_flashcards(text, key_concepts)

def generate_basic_flashcards(text, key_concepts):
    """Generate basic flashcards without AI."""
    
    flashcards = []
    concepts = key_concepts.split(',')[:5]
    
    for concept in concepts:
        concept = concept.strip()
        if concept:
            flashcards.append({
                'term': concept,
                'definition': f"Definition related to {concept} from the content."
            })
    
    return flashcards

def generate_questions(text, key_concepts):
    """Generate questions from content."""
    
    try:
        if openai.api_key:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Generate 3 multiple choice questions in JSON format with 'question', 'answer', and 'type' fields."},
                    {"role": "user", "content": f"Content: {text[:2000]}\nKey concepts: {key_concepts}"}
                ],
                max_tokens=300
            )
            
            try:
                questions_data = json.loads(response.choices[0].message.content)
                return questions_data if isinstance(questions_data, list) else []
            except:
                pass
        
        # Fallback to basic question generation
        return generate_basic_questions(text, key_concepts)
        
    except Exception as e:
        return generate_basic_questions(text, key_concepts)

def generate_basic_questions(text, key_concepts):
    """Generate basic questions without AI."""
    
    questions = []
    concepts = key_concepts.split(',')[:3]
    
    for concept in concepts:
        concept = concept.strip()
        if concept:
            questions.append({
                'question': f"What is {concept}?",
                'answer': f"Answer related to {concept} from the content.",
                'type': 'multiple_choice'
            })
    
    return questions

def get_ai_response(message, context, user):
    """Get AI response for chat."""
    
    try:
        if OPENAI_API_KEY:
            # Prefer modern client
            if openai_client is not None:
                try:
                    resp = openai_client.responses.create(
                        model="gpt-4o-mini",
                        input=[
                            {
                                "role": "system",
                                "content": f"You are an AI tutor helping a {getattr(user.age_group, 'value', 'student')} student with {context} questions. Provide helpful, educational responses."
                            },
                            {"role": "user", "content": message}
                        ],
                        max_output_tokens=500,
                    )
                    text = getattr(resp, "output_text", None)
                    if text:
                        return text
                except Exception:
                    pass
            # Legacy fallback
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": f"You are an AI tutor helping a {getattr(user.age_group, 'value', 'student')} student with {context} questions. Provide helpful, educational responses."},
                    {"role": "user", "content": message}
                ],
                max_tokens=500
            )
            return response.choices[0].message.content
        else:
            return "AI chat is currently unavailable. Please try again later."
            
    except Exception as e:
        return f"Error getting AI response: {str(e)}"



@ai_bp.route('/process-text', methods=['POST'])
@login_required
def process_text():
    """Process text content for AI content generation."""
    text_content = request.form.get('text_content')
    subject = request.form.get('subject', 'general')
    
    if not text_content:
        flash('No text content provided', 'error')
        return redirect(url_for('ai_services.upload_document'))
    
    try:
        # Generate AI content from text
        ai_content = generate_ai_content(text_content, subject)
        
        # Save to database
        document = Document(
            filename=f'text_input_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt',
            subject=subject,
            user_id=current_user.id,
            content_length=len(text_content),
            file_type='txt'
        )
        db.session.add(document)
        db.session.flush()
        
        # Save extracted content
        content_obj = Content(
            document_id=document.id,
            raw_content=text_content,
            ai_generated_notes=ai_content['notes'],
            ai_generated_summary=ai_content['summary'],
            key_concepts=ai_content['key_concepts'],
            user_id=current_user.id
        )
        db.session.add(content_obj)
        
        # Generate flashcards
        flashcards = generate_flashcards(text_content, ai_content['key_concepts'])
        for flashcard_data in flashcards:
            flashcard = Flashcard(
                term=flashcard_data['term'],
                definition=flashcard_data['definition'],
                context=flashcard_data.get('context'),
                example=flashcard_data.get('example'),
                lesson_id=1  # Default lesson ID
            )
            db.session.add(flashcard)
        
        db.session.commit()
        flash('Text content processed successfully!', 'success')
        return redirect(url_for('ai_services.view_content', content_id=content_obj.id))
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error processing text: {str(e)}', 'error')
        return redirect(url_for('ai_services.upload_document'))

@ai_bp.route('/process-webpage', methods=['POST'])
@login_required
def process_webpage():
    """Process webpage content for AI content generation."""
    webpage_url = request.form.get('webpage_url')
    subject = request.form.get('subject', 'general')
    
    if not webpage_url:
        flash('No URL provided', 'error')
        return redirect(url_for('ai_services.upload_document'))
    
    try:
        # Extract content from webpage
        response = requests.get(webpage_url)
        response.raise_for_status()
        
        # Simple HTML parsing (you might want to use BeautifulSoup for better parsing)
        content = re.sub(r'<[^>]+>', '', response.text)
        content = re.sub(r'\s+', ' ', content).strip()
        
        # Generate AI content
        ai_content = generate_ai_content(content, subject)
        
        # Save to database
        document = Document(
            filename=f'webpage_{datetime.now().strftime("%Y%m%d_%H%M%S")}.html',
            subject=subject,
            user_id=current_user.id,
            content_length=len(content),
            file_type='html',
            source_url=webpage_url
        )
        db.session.add(document)
        db.session.flush()
        
        # Save extracted content
        content_obj = Content(
            document_id=document.id,
            raw_content=content,
            ai_generated_notes=ai_content['notes'],
            ai_generated_summary=ai_content['summary'],
            key_concepts=ai_content['key_concepts'],
            user_id=current_user.id
        )
        db.session.add(content_obj)
        
        db.session.commit()
        flash('Text content processed successfully!', 'success')
        return redirect(url_for('ai_services.view_content', content_id=content_obj.id))
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error processing webpage: {str(e)}', 'error')
        return redirect(url_for('ai_services.upload_document'))

@ai_bp.route('/process-youtube', methods=['POST'])
@login_required
def process_youtube():
    """Process YouTube video URL for content extraction."""
    try:
        youtube_url = request.form.get('youtube_url')
        subject = request.form.get('subject')
        description = request.form.get('description', '')
        
        # Check which features to generate
        generate_notes = 'generate_notes' in request.form
        generate_summary = 'generate_summary' in request.form
        generate_flashcards = 'generate_flashcards' in request.form
        generate_questions = 'generate_questions' in request.form
        
        if not youtube_url or not subject:
            flash('YouTube URL and subject are required.', 'error')
            return redirect(url_for('ai_services.upload_document'))
        
        # Extract video information
        try:
            yt = YouTube(youtube_url)
            video_title = yt.title
            video_description = yt.description
            video_duration = yt.length
        except Exception as e:
            flash(f'Error extracting video information: {str(e)}', 'error')
            return redirect(url_for('ai_services.upload_document'))
        
        # Create document record
        document = Document(
            filename=f"youtube_{video_title[:50]}.txt",
            original_filename=f"youtube_{video_title[:50]}.txt",
            file_path="youtube_content",
            file_size=len(video_description),
            content_type="text/plain",
            subject=subject,
            description=description,
            user_id=current_user.id,
            source_type="youtube",
            source_url=youtube_url
        )
        db.session.add(document)
        db.session.flush()
        
        # Create content record
        content = Content(
            document_id=document.id,
            user_id=current_user.id,
            raw_content=f"Title: {video_title}\n\nDescription: {video_description}",
            content_type="youtube",
            processing_status="completed"
        )
        db.session.add(content)
        db.session.flush()
        
        # Generate AI content based on selected features
        if generate_notes or generate_summary or generate_flashcards or generate_questions:
            # Placeholder for AI processing
            content.ai_notes = f"AI-generated notes for: {video_title}"
            content.ai_summary = f"Summary: {video_description[:200]}..."
            
            if generate_flashcards:
                # Create sample flashcards
                flashcard1 = Flashcard(
                    content_id=content.id,
                    term="Video Topic",
                    definition=f"Main topic: {video_title}",
                    difficulty_level="beginner"
                )
                flashcard2 = Flashcard(
                    content_id=content.id,
                    term="Key Concept",
                    definition="Important concept from the video",
                    difficulty_level="intermediate"
                )
                db.session.add(flashcard1)
                db.session.add(flashcard2)
            
            if generate_questions:
                # Create sample questions
                question1 = Question(
                    content_id=content.id,
                    question_text=f"What is the main topic of '{video_title}'?",
                    question_type="multiple_choice",
                    options=["Option A", "Option B", "Option C", "Option D"],
                    correct_answer=0,
                    difficulty_level="beginner"
                )
                db.session.add(question1)
        
        db.session.commit()
        
        flash('YouTube video processed successfully!', 'success')
        return redirect(url_for('ai_services.view_content', content_id=content.id))
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error processing YouTube video: {str(e)}', 'error')
        return redirect(url_for('ai_services.upload_document'))

@ai_bp.route('/web-search', methods=['GET', 'POST'])
@login_required
def web_search():
    """Web search functionality to find and process educational content."""
    if request.method == 'POST':
        try:
            search_query = request.form.get('search_query')
            subject = request.form.get('subject', 'general')
            max_results = int(request.form.get('max_results', 5))
            
            if not search_query:
                flash('Search query is required.', 'error')
                return redirect(url_for('ai_services.web_search'))
            
            # Prefer Google Custom Search API if keys are available
            search_results = perform_web_search(search_query, max_results)
            
            # Process each result
            processed_results = []
            for result in search_results:
                try:
                    url = result['url']
                    # If result is a document (pdf/docx), download and process
                    if url.lower().endswith(('.pdf', '.docx')):
                        file_resp = requests.get(url, timeout=15)
                        if file_resp.status_code == 200:
                            filename = secure_filename(url.split('/')[-1])
                            temp_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
                            with open(temp_path, 'wb') as f:
                                f.write(file_resp.content)
                            extracted_text = extract_document_content(temp_path, filename)
                            ai_content = generate_ai_content(extracted_text, subject)
                            document = Document(
                                filename=filename,
                                subject=subject,
                                user_id=current_user.id,
                                content_length=len(extracted_text),
                                file_type=filename.rsplit('.', 1)[1].lower(),
                                source_url=url
                            )
                            db.session.add(document)
                            db.session.flush()
                            content_obj = Content(
                                document_id=document.id,
                                raw_content=extracted_text,
                                ai_generated_notes=ai_content['notes'],
                                ai_generated_summary=ai_content['summary'],
                                key_concepts=ai_content['key_concepts'],
                                user_id=current_user.id
                            )
                            db.session.add(content_obj)
                            try:
                                lesson = create_lesson_from_content(
                                    title=os.path.splitext(filename)[0],
                                    subject=subject,
                                    summary=ai_content['summary'],
                                    user_id=current_user.id,
                                    age_group=current_user.age_group if hasattr(current_user, 'age_group') else AgeGroup.TEENS,
                                    format_type=filename.rsplit('.', 1)[1].lower()
                                )
                                db.session.add(lesson)
                            except Exception as _:
                                pass
                            processed_results.append({
                                'title': result['title'],
                                'url': url,
                                'content_id': None,
                                'snippet': ai_content['summary']
                            })
                            try:
                                os.remove(temp_path)
                            except Exception:
                                pass
                    else:
                        # Otherwise treat as webpage
                        content = extract_webpage_content(url)
                        
                        if content and len(content) > 100:
                            document = Document(
                                filename=f"web_{result['title'][:50]}.html",
                                subject=subject,
                                user_id=current_user.id,
                                content_length=len(content),
                                file_type='html',
                                source_url=url
                            )
                            db.session.add(document)
                            db.session.flush()
                            content_obj = Content(
                                document_id=document.id,
                                raw_content=content,
                                ai_generated_notes='',
                                ai_generated_summary='',
                                key_concepts=[],
                                user_id=current_user.id
                            )
                            db.session.add(content_obj)
                            processed_results.append({
                                'title': result['title'],
                                'url': url,
                                'content_id': content_obj.id,
                                'snippet': content[:200] + "..." if len(content) > 200 else content
                            })
                    
                    if content and len(content) > 100:  # Only process if we got meaningful content
                        # Create document record
                        document = Document(
                            filename=f"web_{result['title'][:50]}.html",
                            original_filename=f"web_{result['title'][:50]}.html",
                            file_path="web_content",
                            file_size=len(content),
                            content_type="text/html",
                            subject=subject,
                            description=f"Web search result: {search_query}",
                            user_id=current_user.id,
                            source_type="web",
                            source_url=result['url']
                        )
                        db.session.add(document)
                        db.session.flush()
                        
                        # Create content record
                        content_obj = Content(
                            document_id=document.id,
                            user_id=current_user.id,
                            raw_content=content,
                            content_type="web",
                            processing_status="completed"
                        )
                        db.session.add(content_obj)
                        db.session.flush()
                        
                        processed_results.append({
                            'title': result['title'],
                            'url': result['url'],
                            'content_id': content_obj.id,
                            'snippet': content[:200] + "..." if len(content) > 200 else content
                        })
                        
                except Exception as e:
                    print(f"Error processing {result['url']}: {str(e)}")
                    continue
            
            db.session.commit()
            
            flash(f'Successfully processed {len(processed_results)} web results!', 'success')
            return render_template('ai_services/web_search_results.html', 
                                 results=processed_results, 
                                 query=search_query)
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error performing web search: {str(e)}', 'error')
            return redirect(url_for('ai_services.web_search'))
    
    return render_template('ai_services/web_search.html')

# Route aliases for convenience
@ai_bp.route('/google', methods=['GET', 'POST'])
@ai_bp.route('/search-docs', methods=['GET', 'POST'])
@login_required
def web_search_alias():
    return web_search()

def perform_web_search(query, max_results=5):
    """Perform web search via SerpAPI or Google CSE if configured, else fallback."""
    # Prefer SerpAPI if provided (doesn't require CSE ID)
    if SERPAPI_KEY:
        try:
            params = {
                'engine': 'google',
                'q': query,
                'num': max_results,
                'api_key': SERPAPI_KEY
            }
            resp = requests.get('https://serpapi.com/search.json', params=params, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            results = []
            for item in data.get('organic_results', [])[:max_results]:
                results.append({
                    'title': item.get('title', ''),
                    'url': item.get('link', ''),
                    'snippet': item.get('snippet', '')
                })
            if results:
                return results
        except Exception:
            pass
    api_key = os.environ.get('GOOGLE_API_KEY')
    cse_id = os.environ.get('GOOGLE_CSE_ID')
    if api_key and cse_id:
        try:
            params = {
                'key': api_key,
                'cx': cse_id,
                'q': query,
                'num': max_results
            }
            resp = requests.get('https://www.googleapis.com/customsearch/v1', params=params, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            items = data.get('items', [])
            results = []
            for item in items:
                results.append({
                    'title': item.get('title', ''),
                    'url': item.get('link', ''),
                    'snippet': item.get('snippet', '')
                })
            return results
        except Exception as _:
            pass
    # Fallback
    return [
        {
            'title': f'Educational Resource: {query}',
            'url': f'https://example.com/education/{query.replace(" ", "-")}',
            'snippet': f'Comprehensive educational content about {query}...'
        }
    ]

def extract_webpage_content(url):
    """Extract text content from a webpage."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        
        # Simple HTML parsing (you might want to use BeautifulSoup for better parsing)
        content = re.sub(r'<[^>]+>', '', response.text)
        content = re.sub(r'\s+', ' ', content).strip()
        
        return content
    except Exception as e:
        print(f"Error extracting content from {url}: {str(e)}")
        return None

@ai_bp.route('/view-content/<int:content_id>')
def view_content(content_id):
    """View processed AI content."""
    content = Content.query.get_or_404(content_id)
    return render_template('ai_services/view_content.html', content=content)


# Internal helper to create a Lesson row from extracted content
def create_lesson_from_content(title: str, subject: str, summary: str, user_id: int, age_group: AgeGroup, format_type: str) -> Lesson:
    format_map = {
        'pdf': ContentFormat.PDF,
        'docx': ContentFormat.DOCX,
        'txt': ContentFormat.TEXT,
        'md': ContentFormat.TEXT,
        'ppt': ContentFormat.TEXT,
        'pptx': ContentFormat.TEXT,
        'html': ContentFormat.TEXT
    }
    ft = format_map.get(format_type.lower(), ContentFormat.TEXT)
    return Lesson(
        title=title[:200],
        description=summary[:500] if summary else None,
        topic=subject.title(),
        subject=subject,
        teacher_id=user_id,
        age_group_target=age_group,
        format_type=ft,
        ai_summary=summary,
        is_published=True,
        tags=[]
    )
